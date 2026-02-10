import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from datetime import datetime
import re  # 新增这行，用于正则表达式处理
import numpy as np  # 必须导入numpy，否则np.nan会报错
import os

class ReportGenerator:
    def __init__(self):
        """初始化报告生成器"""
        self.logger = None  # 可外接日志函数
        self.success = False
        self.error_msg = ""
        self.output_file = ""
        # 原字典
        self.program_replacements = {
            '朝闻天下': '朝闻TX', '生活圈': '生活Q', '无所畏惧之永': '无所畏惧之永',
            '新闻30分': 'XW30分', '今日说法': 'JR说法', '人生之路': '人生之路',
            '第1动画乐园': '第1动画LY', '农耕探文明': '农耕探文明', '新闻联播': 'XW联播',
            '焦点访谈': '焦点FT', '乌蒙深处': '乌蒙深处', '百年守护': '百年守护',
            '中华考工记': '中华考工记', '晚间新闻': '晚间XW',
        }
        # 方法1：字典推导式（简洁高效，推荐）
        self.program_replacements = {value: key for key, value in self.program_replacements.items()}

    def set_logger(self, logger_func):
        """设置日志输出函数（对接main程序的日志系统）"""
        self.logger = logger_func

    def _log(self, message, level="info"):
        """内部日志函数"""
        if self.logger:
            self.logger(f"[报告生成] {message}")
        else:
            if level == "error":
                print(f"❌ {message}")
            elif level == "warning":
                print(f"⚠️ {message}")
            else:
                print(f"ℹ️ {message}")

    def get_time_slot(self, time_str):
        """
        根据输入的时间字符串（如"17:00"）返回对应的档期名称

        参数:
            time_str (str): 时间字符串，格式为"HH:MM"

        返回:
            str: 对应的档期名称，不在范围内则返回空字符串
        """
        # 定义各档期的时间范围（起始分钟, 结束分钟, 档期名称）
        time_slots = [
            (5 * 60, 8 * 60 + 54, "早间档"),  # 5:00-8:54
            (8 * 60 + 55, 11 * 60 + 29, "上午档"),  # 8:55-11:29
            (11 * 60 + 30, 13 * 60 + 29, "中午档"),  # 11:30-13:29
            (13 * 60 + 30, 16 * 60 + 54, "下午档"),  # 13:30-16:54
            (16 * 60 + 55, 18 * 60 + 55, "傍晚档"),  # 16:55-18:55
            (19 * 60 + 58, 21 * 60 + 55, "黄金档"),  # 19:58-21:55
            (22 * 60 + 10, 22 * 60 + 55, "次黄档"),  # 22:10-22:55
            (22 * 60 + 56, 23 * 60 + 59, "晚间档")  # 22:56-23:59
        ]

        try:
            # 将时间字符串转换为总分钟数（便于区间比较）
            hour, minute = map(int, time_str.split(":"))
            total_minutes = hour * 60 + minute

            # 遍历档期列表，判断时间所属区间
            for start, end, slot_name in time_slots:
                if start <= total_minutes <= end:
                    return slot_name

            # 不在任何区间内，返回空字符串
            return ""

        except (ValueError, AttributeError):
            # 处理输入格式错误的情况（如"17:61"、"abc"等），返回空字符串
            return ""

    # 字符替换函数（示例）
    def replace_chars(self, text, replacements):
        if pd.isna(text):
            return text
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text
    """报告生成器类 - 可被外部模块导入并调用"""
    def clear_table_data(self, table):
        """清空表格除表头外的所有数据行"""
        for row in table.rows[1:]:
            table._element.remove(row._element)

    def add_table_border(self, table):
        """为Word表格添加普通黑色细线边框（兼容所有版本）"""
        border_xml = f"""
            <w:tblBorders {nsdecls('w')}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            </w:tblBorders>
        """
        tbl_borders = parse_xml(border_xml)
        table._element.tblPr.append(tbl_borders)


    def set_cell_format(self, cell, is_header=False):
        """设置单元格格式：表头可换行，内容单元格不换行+水平居中"""
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.space_before = paragraph.space_after = 0
            # 内容单元格禁用换行（表头保持默认可换行）
            if not is_header:
                # 添加不换行XML属性
                no_wrap_xml = """
                    <w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>
                """
                no_wrap = parse_xml(no_wrap_xml)
                if cell._element.xpath('.//w:noWrap'):
                    cell._element.xpath('.//w:noWrap')[0].getparent().replace(no_wrap)
                else:
                    cell._element.get_or_add_tcPr().append(no_wrap)


    def auto_adjust_table_columns(self, table):
        """表格列宽自适应（兼容所有版本）"""
        table.autofit = True
        # 移除单元格固定列宽
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcW_nodes = tcPr.xpath('.//w:tcW')
                for tcW in tcW_nodes:
                    tcPr.remove(tcW)
        # 设置表格自动布局
        tbl_layout_xml = """
            <w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="auto"/>
        """
        tbl_layout = parse_xml(tbl_layout_xml)
        existing_layout = table._element.xpath('.//w:tblLayout')
        if existing_layout:
            parent = existing_layout[0].getparent()
            parent.remove(existing_layout[0])
            parent.append(tbl_layout)
        else:
            table._element.tblPr.append(tbl_layout)


    def set_cell_shading(self, cell, color_hex):
        """兼容版本的单元格底色设置"""
        shading_xml = f"""
            <w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                  w:fill="{color_hex}"
                  w:val="clear"/>
        """
        shading = parse_xml(shading_xml)
        existing_shd = cell._element.xpath('.//w:shd')
        if existing_shd:
            parent = existing_shd[0].getparent()
            parent.remove(existing_shd[0])
            parent.append(shading)
        else:
            cell._element.get_or_add_tcPr().append(shading)


    def format_table1_row(self, row):
        """表格1：综合频道行设置（黄底+红字+加粗）"""
        for cell in row.cells:
            self.set_cell_shading(cell, "FFFF00")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)


    def format_table2_cell(self, cell, value, col_name):
        """表格2：条件标红"""
        if col_name == "较前一日变化幅度" and value != "/":
            try:
                num = float(value.replace("%", ""))
                if num > 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)
            except:
                pass
        elif col_name == "排名变化" and "↑" in str(value):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 0, 0)
        elif col_name in ["排名1", "排名2", "排名3"] and str(value).strip() == "综合":
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 0, 0)


    def format_table3_cell(self, cell, value, col_name):
        """表格3：条件标红"""
        if col_name == "较前一日变化幅度" and value != "/":
            try:
                num = float(value.replace("%", ""))
                if num > 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)
            except:
                pass
        elif col_name in ["排名1", "排名2", "排名3"] and str(value).strip() == "综合":
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 0, 0)


    def replace_report_title(self, doc, target_date=None):
        """替换标题（保持格式不变）"""
        if target_date is None:
            target_date = datetime.now()
        elif isinstance(target_date, str):
            target_date = datetime.strptime(target_date, "%Y-%m-%d")

        new_title_text = self.GetTitleText(target_date)

        if doc.paragraphs:
            title_para = doc.paragraphs[0]
            original_runs = list(title_para.runs)
            title_para.clear()

            for run in original_runs:
                new_run = title_para.add_run(new_title_text)
                new_run.font.size = run.font.size
                new_run.font.bold = run.font.bold
                new_run.font.italic = run.font.italic
                new_run.font.color.rgb = run.font.color.rgb if run.font.color.rgb else None
                new_run.font.name = run.font.name
                break

            title_para.alignment = title_para.alignment
            print(f"📝 标题已更新为：{new_title_text}（字号保持不变）")
        else:
            print("⚠️  未找到标题段落")

    def GetTitleText(self, target_date):
        year = target_date.year
        month = target_date.month
        day = target_date.day
        week_map = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]
        weekday = week_map[target_date.weekday()]
        new_title_text = f"日报（{year}年 {month} 月 {day}日 {weekday}）"
        return new_title_text

    def GetReportName(self, target_date):
        year = target_date.year
        month = target_date.month
        day = target_date.day
        week_map = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"]
        weekday = week_map[target_date.weekday()]
        # 【日报】xxxx年xx月xx日（周xx）
        report_name = f"【日报】{year}年{month}月{day}日（{weekday}）.docx"
        return report_name

    def fill_competitiveness_template(self, doc, share_df, change_df):
        """填充竞争力情况模板文字"""
        zonghe_data = share_df[share_df["频道"] == "综合"].iloc[0] if "综合" in share_df["频道"].values else None
        if zonghe_data is None:
            raise ValueError("未找到综合频道数据")

        rate = zonghe_data["收视率%"]
        rate_change = zonghe_data["收视率%较前一日变化幅度"]
        share = zonghe_data["收视份额%"]
        share_change = zonghe_data["收视份额%较前一日变化幅度"]
        rank = zonghe_data["排名"]
        rank_change = zonghe_data["排名较前一日变化"]

        def get_trend(change_val):
            if change_val == "持平" or change_val == "0%" or change_val == "0":
                return "持平"
            elif change_val.startswith("-"):
                return "下降"
            elif change_val.startswith("+") or (
                    change_val.replace("%", "").isdigit() and float(change_val.replace("%", "")) > 0):
                return "上升"
            else:
                return "持平"

        def get_change_num(change_val, trend):
            if trend == "持平" or change_val in ("持平", "/", "0%", "0"):
                return ""
            return change_val.replace("↑", "").replace("↓", "").replace("%", "").replace("-", "")

        rate_trend = get_trend(rate_change)
        share_trend = get_trend(share_change)
        rank_trend = get_trend(rank_change)

        rate_change_num = get_change_num(rate_change, rate_trend)
        share_change_num = get_change_num(share_change, share_trend)
        rank_change_num = get_change_num(rank_change, rank_trend)

        top1_count = len(change_df[change_df["全国排名"] == "1"])
        total_program_count = len(change_df)

        # 收视率提升最大节目
        change_df_valid = change_df[change_df["较前一日变化幅度"] != "/"].copy()
        if not change_df_valid.empty:
            change_df_valid["change_num"] = change_df_valid["较前一日变化幅度"].apply(
                lambda x: float(x.replace("%", "")) if x.replace("%", "").lstrip("-").isdigit() else -float("inf")
            )
            increase_df = change_df_valid[change_df_valid["change_num"] > 0]
            max_rate_program = increase_df.loc[increase_df["change_num"].idxmax()] if not increase_df.empty else \
            change_df_valid.loc[change_df_valid["change_num"].idxmax()]
            max_rate_time = max_rate_program["播出时间"]
            max_rate_name = max_rate_program["名称"]
            max_rate_increase = max_rate_program["较前一日变化幅度"]
        else:
            max_rate_time = "异常"
            max_rate_name = "异常"
            max_rate_increase = "异常"

        # 排名提升最大节目
        rank_change_valid = change_df[change_df["排名变化"].str.contains("↑", na=False)].copy()
        if not rank_change_valid.empty:
            rank_change_valid["rank_change_num"] = rank_change_valid["排名变化"].apply(
                lambda x: int(x.replace("↑", "")) if x.replace("↑", "").isdigit() else 0
            )
            max_rank_program = rank_change_valid.loc[rank_change_valid["rank_change_num"].idxmax()]
            max_rank_time = max_rank_program["播出时间"]
            max_rank_name = max_rank_program["名称"]
            max_rank_increase = max_rank_program["排名变化"].replace("↑", "")
        else:
            max_rank_time = "异常"
            max_rank_name = "异常"
            max_rank_increase = "异常"

        # 构建文本
        rate_part = f"综合频道收视率{rate}%，较前一日{rate_trend}" + (f"{rate_change_num}%" if rate_change_num else "") + "；"
        share_part = f"收视份额{share}%，较前一日{share_trend}" + (f"{share_change_num}%" if share_change_num else "") + "；"
        rank_part = f"排名全国第{rank}位，与前一日{rank_trend}" + (f"{rank_change_num}位" if rank_change_num else "")
        template_text1 = rate_part + share_part + rank_part + "。"

        max_rate_time = self.get_time_slot(max_rate_time)
        max_rank_time = self.get_time_slot(max_rank_time)
        max_rate_name = self.replace_chars(max_rate_name, self.program_replacements)
        max_rank_name = self.replace_chars(max_rank_name, self.program_replacements)
        template_text2 = f"""全天{total_program_count}档节目中共{top1_count}档在全国排名第一。{max_rate_time}《{max_rate_name}》收视率较前一日提升{max_rate_increase}，收视率提升幅度最大；{max_rank_time}《{max_rank_name}》全国排名提升{max_rank_increase}位，排名提升幅度最大。"""
        if max_rate_name == "异常" and max_rate_increase == "异常":
            template_text2 = f"""全天{total_program_count}档节目中共{top1_count}档在全国排名第一。{max_rate_time}《{max_rate_name}》收视率较前一日提升{max_rate_increase}，收视率提升幅度最大。"""
        # 替换模板段落
        for idx, para in enumerate(doc.paragraphs):
            if "综合频道收视率【xx】%" in para.text:
                original_runs = list(para.runs)
                para.clear()
                for run in original_runs:
                    new_run = para.add_run(template_text1)
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
                    new_run.font.bold = run.font.bold
                    new_run.font.color.rgb = run.font.color.rgb if run.font.color.rgb else None
                    break
                print("📊 第一段模板填充完成")
                break

        for idx, para in enumerate(doc.paragraphs):
            if "全天【xx】档节目中共" in para.text:
                original_runs = list(para.runs)
                para.clear()
                for run in original_runs:
                    new_run = para.add_run(template_text2)
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
                    new_run.font.bold = run.font.bold
                    new_run.font.color.rgb = run.font.color.rgb if run.font.color.rgb else None
                    break
                print("📊 第二段模板填充完成")
                break


    def insert_data_to_table(self, table, data_df, target_cols, table_type=1):
        """插入数据+格式设置（区分表头/内容单元格）"""
        data_df = data_df.reindex(columns=target_cols, fill_value="")

        # 插入数据行（内容单元格）
        for row_idx, (_, row) in enumerate(data_df.iterrows()):
            new_row = table.add_row().cells
            for col_idx, (col_name, value) in enumerate(zip(target_cols, row)):
                if col_idx < len(new_row):
                    cell = new_row[col_idx]
                    # 处理空值
                    if pd.isna(value) or value is None or str(value).strip() in ("nan", "None", ""):
                        cell.text = "/"
                    else:
                        cell.text = str(value).strip()
                    # 内容单元格：不换行+居中
                    self.set_cell_format(cell, is_header=False)
                    # 条件格式
                    if table_type == 2:
                        self.format_table2_cell(cell, cell.text, col_name)
                    elif table_type == 3:
                        self.format_table3_cell(cell, cell.text, col_name)

        # 表头格式（可换行+居中）
        for cell in table.rows[0].cells:
            self.set_cell_format(cell, is_header=True)

        # 表格1综合频道行高亮
        if table_type == 1:
            for row in table.rows[1:]:
                channel_cell = row.cells[1]
                if channel_cell.text.strip() == "综合":
                    self.format_table1_row(row)

        # 边框+自适应
        self.add_table_border(table)
        if table_type == 2:
            self.auto_adjust_table_columns(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER


    def generate_report(self, excel_path, word_template_path):
        TABLE1_COLS = [
            "排名", "频道", "收视率%", "收视率%较前一日变化幅度",
            "收视份额%", "收视份额%较前一日变化幅度", "排名较前一日变化"
        ]
        TABLE2_COLS = [
            "播出时间", "名称", "收视率%", "较前一日变化幅度",
            "全国排名", "排名变化", "排名1", "排名2", "排名3"
        ]
        TABLE3_COLS = [
            "播出时间", "名称", "收视率", "较前一日变化幅度",
            "全国排名", "排名变化", "排名1", "排名2", "排名3"
        ]

        try:
            # 读取Excel数据
            share_df = pd.read_excel(
                excel_path,
                sheet_name="份额全国排名",
                usecols="A:F",
                nrows=10,
                dtype=str
            )
            if "排名" not in share_df.columns:
                share_df.insert(0, "排名", list(range(1, 11)))
            share_df.columns = TABLE1_COLS[:len(share_df.columns)]

            change_df = pd.read_excel(
                excel_path,
                sheet_name="一套变化情况",
                usecols="A:J",
                dtype=str
            )
            change_df = change_df[change_df.iloc[:, 9] == "本日"].reset_index(drop=True)
            # 删除日期列
            date_col_candidates = [col for col in change_df.columns if "日期" in str(col).strip()]
            if date_col_candidates:
                change_df = change_df.drop(columns=date_col_candidates[0], errors="ignore")
            elif len(change_df.columns) >= 10:
                change_df = change_df.drop(columns=change_df.columns[9], errors="ignore")
            # 列名映射
            col_mapping = {
                "开始时间": "播出时间",
                "节目名称": "名称",
                "中央电视台综合频道": "收视率%",
                "收视变化": "较前一日变化幅度",
                "全国排名": "全国排名",
                "排名变化": "排名变化",
                "排名1": "排名1",
                "排名2": "排名2",
                "排名3": "排名3"
            }
            change_df = change_df.rename(columns=lambda x: col_mapping.get(str(x).strip(), str(x).strip()))
            change_df = change_df[TABLE2_COLS].fillna("")

            drama_df = pd.read_excel(
                excel_path,
                sheet_name="电视剧频道黄金时段电视剧",
                usecols="A:J",
                dtype=str
            )
            drama_df = drama_df[drama_df["日期"] == "本日"].reset_index(drop=True)
            drama_col_mapping = {
                "开始时间": "播出时间",
                "节目名称": "名称",
                "中央台八套": "收视率",
                "收视变化": "较前一日变化幅度",
                "全国排名": "全国排名",
                "排名变化": "排名变化",
                "排名1": "排名1",
                "排名2": "排名2",
                "排名3": "排名3"
            }
            drama_df = drama_df.rename(columns=lambda x: drama_col_mapping.get(str(x).strip(), str(x).strip()))
            drama_df = drama_df[TABLE3_COLS].fillna("")

            # 操作Word模板
            doc = Document(word_template_path)
            if len(doc.tables) < 3:
                raise ValueError("模板需包含3个表格")

            totalrate_df = pd.read_excel(
                excel_path,
                sheet_name="分钟总收视",
                usecols="A:B",
                dtype=str
            )
            report_date = totalrate_df.iloc[0, 1]
            report_date_fixed = datetime.strptime(report_date, "%Y/%m/%d").strftime("%Y-%m-%d")
            print(f"成功提取日期值：{report_date_fixed}")
            # 替换标题
            self.replace_report_title(doc, report_date_fixed)

            # 填充模板文字
            self.fill_competitiveness_template(doc, share_df, change_df)

            # 插入表格数据
            table1 = doc.tables[0]
            self.clear_table_data(table1)
            self.insert_data_to_table(table1, share_df, TABLE1_COLS, table_type=1)

            table2 = doc.tables[1]
            self.clear_table_data(table2)
            self.insert_data_to_table(table2, change_df, TABLE2_COLS, table_type=2)

            table3 = doc.tables[2]
            self.clear_table_data(table3)
            self.insert_data_to_table(table3, drama_df, TABLE3_COLS, table_type=3)

            report_date = datetime.strptime(report_date, "%Y/%m/%d")
            excel_dir = os.path.dirname(excel_path)
            output_word_name = self.GetReportName(report_date)
            output_word_path = os.path.join(excel_dir, output_word_name)
            # 保存文件
            doc.save(output_word_path)
            print(f"🎉 生成成功！文件：{output_word_path}")
            print(f"📊 数据统计：表1={len(share_df)}行 | 表2={len(change_df)}行 | 表3={len(drama_df)}行")
            self.success = True
            return (self.success, "报告生成成功", output_word_path)
        except Exception as e:
            error_msg = f"❌ 报错：{str(e)}"
            import traceback
            traceback.print_exc()
            return (False, error_msg, "")

    def standardize_program_name(self, name):
        """
        标准化节目名称，处理第一动画乐园的特殊情况

        参数:
        name: 原始节目名称

        返回:
        标准化后的节目名称
        """
        if pd.isna(name):
            return name

        name_str = str(name).strip()

        # 处理第一动画乐园的分集格式（如"第1动画乐园:哪吒传奇"、"第一动画乐园-第一集"等）
        animation_patterns = [
            r'第1动画乐园[:-].*',  # 匹配"第1动画乐园:xxx"或"第1动画乐园-xxx"
            r'第一动画乐园[:-].*'  # 匹配"第一动画乐园:xxx"或"第一动画乐园-xxx"
        ]

        for pattern in animation_patterns:
            if re.match(pattern, name_str):
                # 提取主节目名称（取冒号/横杠前面的部分）
                if ':' in name_str:
                    main_name = name_str.split(':')[0].strip()
                elif '-' in name_str:
                    main_name = name_str.split('-')[0].strip()
                else:
                    main_name = name_str

                # 统一为"第1动画乐园"格式
                if '第一动画乐园' in main_name:
                    return '第1动画乐园'
                return main_name

        return name_str

    def merge_animation_records(self, df, name_column='标准化名称'):
        """
        优化版：合并第一动画乐园的多条记录（适配仅含名称+1-150分钟列的场景）
        核心逻辑：
        1. 筛选所有第1动画乐园记录
        2. 对每条记录遍历分钟列，遇到空值（NaN/空字符串/None）立即停止提取
        3. 按原始顺序拼接所有有效分钟数据
        4. 超出150分钟的部分截断，不足的部分填0
        5. 合并后的记录保留在原表格中第一条第1动画乐园的位置
        6. 计算合并后的总时长并返回

        参数:
        df: 分钟曲线数据框（仅含名称+1分钟~150分钟列）
        name_column: 名称列名

        返回:
        tuple: (合并后的数据框, 动画乐园合并后的总时长)
        """
        # 筛选第一动画乐园的记录
        animation_mask = df[name_column] == '第1动画乐园'
        animation_df = df[animation_mask].copy()

        if len(animation_df) <= 1:
            # 没有多条记录，直接返回原数据和原时长
            if len(animation_df) == 1:
                original_duration = df.loc[df[name_column] == '第1动画乐园', '时长[总和]'].iloc[
                    0] if '时长[总和]' in df.columns else 0
                return df, original_duration
            return df, 0

        self._log(f"发现{len(animation_df)}条第1动画乐园记录，开始拼接分钟数据...")

        # ===== 步骤1：获取第一条第1动画乐园记录的位置 =====
        first_animation_idx = df[animation_mask].index[0]  # 第一条动画乐园记录的原始索引
        self._log(f"第一条第1动画乐园记录的原始位置：行索引{first_animation_idx}")

        # ===== 步骤2：提取并排序所有分钟列 =====
        # 匹配"1分钟/2分钟/.../150分钟"格式的列
        minute_pattern = re.compile(r'^(\d+)分钟$')
        minute_columns = []
        for col in df.columns:
            match = minute_pattern.match(str(col).strip())
            if match:
                minute_columns.append((int(match.group(1)), col))

        # 按分钟数升序排列（1分钟→150分钟）
        minute_columns.sort(key=lambda x: x[0])
        minute_col_names = [col for _, col in minute_columns]  # 排好序的分钟列名
        max_minute_count = len(minute_col_names)  # 最大分钟数（150）

        # ===== 步骤3：提取每条记录的有效分钟数据（遇到空值停止）=====
        merged_minutes_list = []  # 存储拼接后的所有有效分钟数据

        # 遍历每条动画乐园记录
        for record_idx, (_, record) in enumerate(animation_df.iterrows()):
            self._log(f"处理第{record_idx + 1}条第1动画乐园记录...")

            # 遍历当前记录的分钟列，提取有效数据
            current_valid_minutes = []
            for min_col in minute_col_names:
                # 获取当前分钟值
                val = record[min_col]

                # 判断是否为空值（包含NaN/空字符串/None/仅空白字符）
                is_empty = False
                if pd.isna(val):
                    is_empty = True
                elif isinstance(val, str) and val.strip() == "":
                    is_empty = True
                elif val is None:
                    is_empty = True

                # 遇到空值立即停止提取当前记录的分钟数据
                if is_empty:
                    self._log(f"  第{min_col}值为空，停止提取当前记录的分钟数据")
                    break

                # 转换为数值类型（非空值）
                try:
                    val_num = float(val)
                except (ValueError, TypeError):
                    # 非数值但非空的情况（如文本），视为有效数据保留
                    val_num = val

                # 有效数据加入列表
                current_valid_minutes.append(val_num)

            # 将当前记录的有效分钟数据加入总列表
            merged_minutes_list.extend(current_valid_minutes)
            self._log(f"  提取到{len(current_valid_minutes)}个有效分钟数据")

        # ===== 步骤4：计算合并后的总时长 =====
        total_duration = len(merged_minutes_list)  # 有效分钟数即为总时长
        self._log(f"第1动画乐园合并后的总时长：{total_duration}分钟")

        # ===== 步骤5：构建合并后的分钟数据 =====
        merged_minutes = {}
        for idx, min_col in enumerate(minute_col_names):
            if idx < len(merged_minutes_list):
                # 有拼接数据的位置填入对应值
                merged_minutes[min_col] = merged_minutes_list[idx]
            else:
                # 超出拼接长度的位置填0
                merged_minutes[min_col] = ''

        # ===== 步骤6：构建合并后的完整记录 =====
        # 初始化合并记录（保留第一条记录的非分钟列数据，仅替换分钟列）
        first_animation_record = df.loc[first_animation_idx].copy()
        merged_record = first_animation_record.to_dict()
        # 更新分钟列数据
        merged_record.update(merged_minutes)
        # 确保名称列正确
        merged_record[name_column] = '第1动画乐园'

        # ===== 步骤7：生成最终数据框（保留原始位置）=====
        # 1. 移除所有第1动画乐园记录
        df_no_animation = df[~animation_mask].copy()

        # 2. 将合并后的记录插入到第一条动画乐园的原始位置
        # 先将数据框拆分为前半部分和后半部分
        df_before = df_no_animation[df_no_animation.index < first_animation_idx].copy()
        df_after = df_no_animation[df_no_animation.index > first_animation_idx].copy()

        # 3. 转换合并记录为DataFrame
        merged_df = pd.DataFrame([merged_record])

        # 4. 拼接所有部分，保持原始顺序
        result_df = pd.concat([
            df_before,
            merged_df,
            df_after
        ], ignore_index=True)

        self._log(
            f"第1动画乐园记录拼接完成：共提取{len(merged_minutes_list)}个有效分钟数据 | 合并后剩余{len(result_df)}条记录")
        self._log(f"合并后的第1动画乐园记录已放置在原始第一条的位置（当前行索引：{len(df_before)}）")

        return result_df, total_duration

    def merge_tv_ratings_data(self, input_file):
        """
        主函数：合并电视收视数据

        参数:
        input_file: 输入Excel文件路径
        output_file: 输出Excel文件路径

        返回:
        success,message,output_file
        """
        success = False
        message = ""
        output_excel_path = ""
        try:
            # 1. 读取数据
            print("正在读取数据...")
            # 读取分钟曲线数据
            try:
                df_curve = pd.read_excel(input_file, sheet_name='分钟曲线')
                print(f"成功读取'分钟曲线'数据：{len(df_curve)}行 × {len(df_curve.columns)}列")
            except Exception as e:
                raise Exception(f"读取'分钟曲线'sheet失败：{str(e)}")

            # 读取分钟总收视数据
            try:
                df_total = pd.read_excel(input_file, sheet_name='分钟总收视')
                print(f"成功读取'分钟总收视'数据：{len(df_total)}行 × {len(df_total.columns)}列")
            except Exception as e:
                raise Exception(f"读取'分钟总收视'sheet失败：{str(e)}")

            # 2. 数据预处理 - 标准化节目名称
            print("\n正在标准化节目名称...")
            # 为两个数据框添加标准化名称列
            df_curve['标准化名称'] = df_curve['名称'].apply(self.standardize_program_name)
            df_total['标准化名称'] = df_total['名称'].apply(self.standardize_program_name)

            # 3. 合并第一动画乐园的多条记录（仅对分钟曲线数据）
            print("\n正在处理第1动画乐园记录...")
            df_curve_processed, animation_duration = self.merge_animation_records(df_curve)

            # 4. 数据合并 - 按标准化名称进行连接
            print("\n正在合并两个sheet的数据...")
            # 准备合并数据，删除重复的名称列
            df_curve_merge = df_curve_processed.drop('名称', axis=1, errors='ignore')
            df_total_merge = df_total.drop('名称', axis=1, errors='ignore')
            # 对表1按"名称"分组，添加组内序号（从0开始计数）
            df_curve_merge['序号'] = df_curve_merge.groupby('标准化名称').cumcount()
            # 对表2做同样操作，保证同名行的序号一一对应
            df_total_merge['序号'] = df_total_merge.groupby('标准化名称').cumcount()
            # 左连接，以分钟曲线数据为主
            df_combined = pd.merge(
                df_curve_merge,
                df_total_merge,
                on=['标准化名称','序号'],
                how='left',
                suffixes=('_curve', '_total')
            )
            df_combined = df_combined.drop_duplicates(keep='first')
            # print(df_combined.to_string())
            # 5. 列名映射 - 按照要求的输出格式
            print("\n正在整理输出格式...")
            # 定义列名映射关系
            column_mapping = {
                '标准化名称': '名称',
                '日期[相同值]': '日期',
                '周日[具体值]': '周日',
                '开始时间[最小]': '开始时间',
                '时长[总和]': '时长',
                '结束时间[最大]': '结束时间',
                '收视率%': '收视率%',
                '市场份额%': '市场份额%',
                '平均忠实度': '平均忠实度'
            }

            # 构建最终列名列表
            # 基础信息列
            base_columns = [
                '名称', '日期', '周日', '开始时间', '时长', '结束时间',
                '收视率%', '市场份额%', '平均忠实度'
            ]

            # 1-150分钟的列
            minute_columns = [f'{i}分钟' for i in range(1, 151)]

            # 完整列名列表
            final_columns = base_columns + minute_columns

            # 6. 数据整理和清理
            # 重命名列
            df_final = df_combined.rename(columns=column_mapping)

            # 确保所有需要的列都存在
            for col in final_columns:
                if col not in df_final.columns:
                    if col in minute_columns:
                        # 分钟列不存在则填充0
                        df_final[col] = 0
                    else:
                        # 基础列不存在则填充空字符串
                        df_final[col] = ''

            # 选择最终需要的列
            df_final = df_final[final_columns].copy()
            # 数据类型转换和清理
            # 数值列处理
            numeric_cols = ['收视率%', '市场份额%', '平均忠实度']

            for col in numeric_cols:
                # 转换为数值类型，无法转换的设为0
                df_final[col] = pd.to_numeric(df_final[col], errors='coerce').fillna(0)
                # 保留4位小数
                df_final[col] = df_final[col].round(4)

            # 日期列处理
            if '日期' in df_final.columns:
                df_final['日期'] = pd.to_datetime(df_final['日期'], errors='coerce').fillna('')
                df_final['日期'] = df_final['日期'].apply(
                    lambda x: f"{x.year}/{x.month}/{x.day}"  # 单数月份/日期不补零
                )
            # 2. 核心：找到"平均忠实度"列的索引位置（增加列存在性检查）
            try:
                # 获取"平均忠实度"列的索引
                col_index = df_final.columns.get_loc('平均忠实度')
                # 3. 在该列后面插入空列（位置=索引+1）
                # 插入名为"新增空列"的空列，值为pd.NA（pandas原生空值，无依赖）
                df_final.insert(
                    loc=col_index + 1,  # 插入位置："平均忠实度"列的下一位
                    column='',  # 自定义空列名，可修改
                    value=pd.NA  # 空值（无需numpy，兼容所有类型）
                )
            except KeyError:
                # 若"平均忠实度"列不存在，给出提示且不报错
                print("警告：DataFrame中不存在'平均忠实度'列，未插入空列")
            report_date = df_final.iloc[0, 1]
            report_date = datetime.strptime(report_date, "%Y/%m/%d").strftime("%Y年%m月%d日")
            print(f"成功提取日期值：{report_date}")
            excel_dir = os.path.dirname(input_file)
            output_excel_name = f"{report_date}栏目收视率及分钟曲线.xlsx"
            output_excel_path = os.path.join(excel_dir, output_excel_name)
            # 7. 保存输出文件
            print(f"\n正在保存输出文件：{output_excel_path}")
            with pd.ExcelWriter(output_excel_path, engine='openpyxl') as writer:
                # 保存合并后的数据
                df_final.to_excel(writer, sheet_name='合并后收视数据', index=False)
            success = True
            return success,"",output_excel_path

        except Exception as e:
            success = False
            message = f"\n❌ 数据合并过程中发生错误：{str(e)}"
            return success,message,output_file

if __name__ == "__main__":
    EXCEL_FILE = "11111.xlsx"
    EXCEL_OUT_FILE = "0101_111.xlsx"
    WORD_TEMPLATE = "报告模板.docx"

    generator = ReportGenerator()
    # success, msg, output_file = generator.generate_report(EXCEL_FILE, WORD_TEMPLATE)
    success, msg, output_file = generator.merge_tv_ratings_data(EXCEL_FILE)